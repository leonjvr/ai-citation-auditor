#!/usr/bin/env python3
"""
Extract text content from .docx files for reference checking
"""

import sys
from docx import Document

def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(docx_path)
        full_text = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        return '\n'.join(full_text)

    except Exception as e:
        return f"ERROR: Could not extract text from {docx_path}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <path_to_docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    text = extract_text_from_docx(docx_path)

    # Ensure proper UTF-8 encoding for output
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())
    print(text)
